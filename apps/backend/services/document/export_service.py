"""
Document export service for multiple formats (PDF, Word, JSON).
"""

import asyncio
import logging
import json
import os
import tempfile
from typing import Dict, Any, List, Optional, Union
from datetime import datetime
from pathlib import Path
from enum import Enum
from pydantic import BaseModel, Field
import base64

# Document generation libraries
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import weasyprint
from jinja2 import Environment, BaseLoader
import markdown

from core.config import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()


class ExportFormat(str, Enum):
    """Supported export formats."""
    PDF = "pdf"
    WORD = "docx"
    JSON = "json"
    HTML = "html"
    MARKDOWN = "markdown"


class ExportOptions(BaseModel):
    """Options for document export."""
    
    format: ExportFormat = Field(..., description="Export format")
    include_metadata: bool = Field(default=True, description="Include document metadata")
    include_wbs: bool = Field(default=True, description="Include Work Breakdown Structure")
    include_estimates: bool = Field(default=True, description="Include resource estimates")
    template_style: str = Field(default="professional", description="Template style")
    page_size: str = Field(default="A4", description="Page size for PDF/Word")
    margins: Dict[str, float] = Field(default_factory=lambda: {"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0})
    header_footer: bool = Field(default=True, description="Include header and footer")
    table_of_contents: bool = Field(default=True, description="Include table of contents")


class ExportResult(BaseModel):
    """Result of document export operation."""
    
    success: bool = Field(..., description="Export success status")
    file_path: Optional[str] = Field(None, description="Path to exported file")
    file_size_bytes: Optional[int] = Field(None, description="File size in bytes")
    export_time_ms: int = Field(..., description="Export time in milliseconds")
    format: ExportFormat = Field(..., description="Export format")
    error_message: Optional[str] = Field(None, description="Error message if failed")
    metadata: Dict[str, Any] = Field(default_factory=dict, description="Export metadata")


class ExportService:
    """Service for exporting documents to various formats."""
    
    def __init__(self):
        self.export_dir = Path(settings.EXPORT_DIR if hasattr(settings, 'EXPORT_DIR') else '/tmp/exports')
        self.export_dir.mkdir(exist_ok=True, parents=True)
        
        # HTML templates for different formats
        self.html_templates = {
            "professional": self._get_professional_template(),
            "minimal": self._get_minimal_template(),
            "corporate": self._get_corporate_template()
        }
        
        # Initialize Jinja2 environment
        self.jinja_env = Environment(loader=BaseLoader())

    async def export_document(
        self,
        document: Any,  # DocumentResponse object
        format: ExportFormat,
        options: Optional[ExportOptions] = None
    ) -> str:
        """Export document to specified format and return file path."""
        start_time = asyncio.get_event_loop().time()
        
        try:
            if options is None:
                options = ExportOptions(format=format)
            
            logger.info(f"Exporting document {document.id} to {format.value}")
            
            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{document.id}_{timestamp}.{format.value}"
            file_path = self.export_dir / filename
            
            # Export based on format
            if format == ExportFormat.JSON:
                await self._export_json(document, file_path, options)
            elif format == ExportFormat.HTML:
                await self._export_html(document, file_path, options)
            elif format == ExportFormat.MARKDOWN:
                await self._export_markdown(document, file_path, options)
            elif format == ExportFormat.PDF:
                await self._export_pdf(document, file_path, options)
            elif format == ExportFormat.WORD:
                await self._export_word(document, file_path, options)
            else:
                raise ValueError(f"Unsupported export format: {format}")
            
            export_time = int((asyncio.get_event_loop().time() - start_time) * 1000)
            file_size = file_path.stat().st_size if file_path.exists() else 0
            
            logger.info(f"Document exported successfully: {file_path} ({file_size} bytes, {export_time}ms)")
            return str(file_path)
            
        except Exception as e:
            logger.error(f"Document export failed: {e}")
            raise

    async def _export_json(
        self,
        document: Any,
        file_path: Path,
        options: ExportOptions
    ) -> None:
        """Export document as JSON."""
        try:
            # Convert document to dictionary
            doc_data = {
                "id": document.id,
                "title": document.title,
                "document_type": document.document_type,
                "content": document.content,
                "created_at": document.created_at.isoformat(),
                "generation_time_ms": document.generation_time_ms
            }
            
            if options.include_metadata:
                doc_data["metadata"] = document.metadata
            
            if options.include_wbs and document.wbs:
                doc_data["wbs"] = self._serialize_wbs(document.wbs)
            
            if options.include_estimates and document.estimates:
                doc_data["estimates"] = self._serialize_estimates(document.estimates)
            
            # Write JSON file
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(doc_data, f, indent=2, ensure_ascii=False, default=str)
                
        except Exception as e:
            logger.error(f"JSON export failed: {e}")
            raise

    async def _export_html(
        self,
        document: Any,
        file_path: Path,
        options: ExportOptions
    ) -> None:
        """Export document as HTML."""
        try:
            # Get template
            template_content = self.html_templates.get(options.template_style, self.html_templates["professional"])
            template = self.jinja_env.from_string(template_content)
            
            # Prepare template data
            template_data = {
                "title": document.title,
                "document_type": document.document_type.replace('_', ' ').title(),
                "content": document.content,
                "metadata": document.metadata if options.include_metadata else {},
                "wbs": document.wbs if options.include_wbs else None,
                "estimates": document.estimates if options.include_estimates else None,
                "generated_at": document.created_at.strftime("%Y-%m-%d %H:%M:%S"),
                "export_options": options.dict()
            }
            
            # Render HTML
            html_content = template.render(**template_data)
            
            # Write HTML file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
                
        except Exception as e:
            logger.error(f"HTML export failed: {e}")
            raise

    async def _export_markdown(
        self,
        document: Any,
        file_path: Path,
        options: ExportOptions
    ) -> None:
        """Export document as Markdown."""
        try:
            markdown_content = self._generate_markdown_content(document, options)
            
            # Write Markdown file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
                
        except Exception as e:
            logger.error(f"Markdown export failed: {e}")
            raise

    async def _export_pdf(
        self,
        document: Any,
        file_path: Path,
        options: ExportOptions
    ) -> None:
        """Export document as PDF using WeasyPrint."""
        try:
            # First generate HTML
            html_content = await self._generate_html_for_pdf(document, options)
            
            # Generate PDF using WeasyPrint
            weasyprint.HTML(string=html_content).write_pdf(
                str(file_path),
                stylesheets=[weasyprint.CSS(string=self._get_pdf_styles(options))]
            )
            
        except Exception as e:
            logger.error(f"PDF export failed: {e}")
            raise

    async def _export_word(
        self,
        document: Any,
        file_path: Path,
        options: ExportOptions
    ) -> None:
        """Export document as Word document using python-docx."""
        try:
            # Create Word document
            doc = Document()
            
            # Add title
            title = doc.add_heading(document.title, 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Add metadata if requested
            if options.include_metadata:
                doc.add_heading('Document Information', level=1)
                self._add_metadata_to_word(doc, document)
            
            # Add table of contents placeholder if requested
            if options.table_of_contents:
                doc.add_heading('Table of Contents', level=1)
                doc.add_paragraph('[Table of Contents will be generated]')
                doc.add_page_break()
            
            # Add content sections
            self._add_content_to_word(doc, document.content)
            
            # Add WBS if requested
            if options.include_wbs and document.wbs:
                doc.add_page_break()
                doc.add_heading('Work Breakdown Structure', level=1)
                self._add_wbs_to_word(doc, document.wbs)
            
            # Add estimates if requested
            if options.include_estimates and document.estimates:
                doc.add_page_break()
                doc.add_heading('Resource Estimates', level=1)
                self._add_estimates_to_word(doc, document.estimates)
            
            # Save document
            doc.save(str(file_path))
            
        except Exception as e:
            logger.error(f"Word export failed: {e}")
            raise

    def _serialize_wbs(self, wbs: Any) -> Dict[str, Any]:
        """Serialize WBS structure for JSON export."""
        try:
            return {
                "project_name": wbs.project_name,
                "total_estimated_hours": wbs.total_estimated_hours,
                "total_estimated_days": wbs.total_estimated_days,
                "phases": [
                    {
                        "id": phase.id,
                        "name": phase.name,
                        "description": phase.description,
                        "order": phase.order,
                        "estimated_duration_days": phase.estimated_duration_days,
                        "tasks": [
                            {
                                "id": task.id,
                                "name": task.name,
                                "description": task.description,
                                "complexity": task.complexity.value,
                                "priority": task.priority.value,
                                "estimated_hours": task.estimated_hours,
                                "dependencies": task.dependencies,
                                "skills_required": task.skills_required,
                                "deliverables": task.deliverables
                            } for task in phase.tasks
                        ]
                    } for phase in wbs.phases
                ],
                "critical_path": wbs.critical_path,
                "milestones": wbs.milestones,
                "resource_summary": wbs.resource_summary,
                "risk_assessment": wbs.risk_assessment
            }
        except Exception as e:
            logger.warning(f"WBS serialization failed: {e}")
            return {}

    def _serialize_estimates(self, estimates: Any) -> Dict[str, Any]:
        """Serialize resource estimates for JSON export."""
        try:
            return {
                "project_name": estimates.project_name,
                "estimation_date": estimates.estimation_date.isoformat(),
                "team_composition": {
                    "total_team_size": estimates.team_composition.total_team_size,
                    "roles": estimates.team_composition.roles,
                    "skill_distribution": estimates.team_composition.skill_distribution,
                    "estimated_monthly_cost": estimates.team_composition.estimated_monthly_cost,
                    "recommended_duration_months": estimates.team_composition.recommended_duration_months
                },
                "cost_estimate": {
                    "human_resources": estimates.cost_estimate.human_resources,
                    "infrastructure": estimates.cost_estimate.infrastructure,
                    "software_licenses": estimates.cost_estimate.software_licenses,
                    "hardware": estimates.cost_estimate.hardware,
                    "external_services": estimates.cost_estimate.external_services,
                    "contingency": estimates.cost_estimate.contingency,
                    "total_cost": estimates.cost_estimate.total_cost,
                    "confidence_level": estimates.cost_estimate.confidence_level
                },
                "timeline_estimate": {
                    "total_duration_days": estimates.timeline_estimate.total_duration_days,
                    "total_duration_months": estimates.timeline_estimate.total_duration_months,
                    "critical_path_duration": estimates.timeline_estimate.critical_path_duration,
                    "buffer_days": estimates.timeline_estimate.buffer_days,
                    "risk_adjusted_duration": estimates.timeline_estimate.risk_adjusted_duration
                },
                "assumptions": estimates.assumptions,
                "risks": estimates.risks,
                "recommendations": estimates.recommendations
            }
        except Exception as e:
            logger.warning(f"Estimates serialization failed: {e}")
            return {}

    def _generate_markdown_content(self, document: Any, options: ExportOptions) -> str:
        """Generate Markdown content from document."""
        lines = []
        
        # Title
        lines.append(f"# {document.title}")
        lines.append("")
        
        # Metadata
        if options.include_metadata:
            lines.append("## Document Information")
            lines.append(f"- **Type**: {document.document_type.replace('_', ' ').title()}")
            lines.append(f"- **Generated**: {document.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
            lines.append(f"- **Generation Time**: {document.generation_time_ms}ms")
            lines.append("")
        
        # Content sections
        content = document.content
        
        if isinstance(content.get("overview"), str):
            lines.append("## Overview")
            lines.append(content["overview"])
            lines.append("")
        
        if content.get("requirements"):
            lines.append("## Requirements")
            requirements = content["requirements"]
            if isinstance(requirements, list):
                for req in requirements:
                    lines.append(f"- {req}")
            else:
                lines.append(str(requirements))
            lines.append("")
        
        if content.get("implementation"):
            lines.append("## Implementation")
            impl = content["implementation"]
            if isinstance(impl, dict):
                if impl.get("approach"):
                    lines.append(f"**Approach**: {impl['approach']}")
                    lines.append("")
                if impl.get("phases"):
                    lines.append("**Phases**:")
                    for phase in impl["phases"]:
                        lines.append(f"- {phase}")
                    lines.append("")
                if impl.get("technologies"):
                    lines.append("**Technologies**:")
                    for tech in impl["technologies"]:
                        lines.append(f"- {tech}")
                    lines.append("")
            else:
                lines.append(str(impl))
                lines.append("")
        
        if content.get("timeline"):
            lines.append("## Timeline")
            timeline = content["timeline"]
            if isinstance(timeline, dict):
                if timeline.get("total_duration"):
                    lines.append(f"**Duration**: {timeline['total_duration']}")
                    lines.append("")
                if timeline.get("milestones"):
                    lines.append("**Milestones**:")
                    for milestone in timeline["milestones"]:
                        if isinstance(milestone, dict):
                            lines.append(f"- **{milestone.get('name', 'Milestone')}** ({milestone.get('date', 'TBD')}): {milestone.get('description', '')}")
                        else:
                            lines.append(f"- {milestone}")
                    lines.append("")
            else:
                lines.append(str(timeline))
                lines.append("")
        
        # WBS section
        if options.include_wbs and document.wbs:
            lines.extend(self._generate_wbs_markdown(document.wbs))
        
        # Estimates section  
        if options.include_estimates and document.estimates:
            lines.extend(self._generate_estimates_markdown(document.estimates))
        
        return "\n".join(lines)

    def _generate_wbs_markdown(self, wbs: Any) -> List[str]:
        """Generate Markdown content for WBS."""
        lines = []
        lines.append("## Work Breakdown Structure")
        lines.append("")
        lines.append(f"**Total Estimated Hours**: {wbs.total_estimated_hours:,.1f}")
        lines.append(f"**Total Estimated Days**: {wbs.total_estimated_days}")
        lines.append("")
        
        for phase in wbs.phases:
            lines.append(f"### {phase.name}")
            lines.append(phase.description)
            lines.append(f"**Duration**: {phase.estimated_duration_days} days")
            lines.append("")
            
            if phase.tasks:
                lines.append("**Tasks**:")
                for task in phase.tasks:
                    lines.append(f"- **{task.name}** ({task.estimated_hours}h, {task.complexity.value})")
                    lines.append(f"  - {task.description}")
                    if task.skills_required:
                        lines.append(f"  - Skills: {', '.join(task.skills_required)}")
                lines.append("")
        
        return lines

    def _generate_estimates_markdown(self, estimates: Any) -> List[str]:
        """Generate Markdown content for resource estimates."""
        lines = []
        lines.append("## Resource Estimates")
        lines.append("")
        
        # Team composition
        lines.append("### Team Composition")
        lines.append(f"**Total Team Size**: {estimates.team_composition.total_team_size}")
        lines.append(f"**Monthly Cost**: ${estimates.team_composition.estimated_monthly_cost:,.0f}")
        lines.append(f"**Duration**: {estimates.team_composition.recommended_duration_months} months")
        lines.append("")
        
        lines.append("**Roles**:")
        for role, count in estimates.team_composition.roles.items():
            lines.append(f"- {role}: {count}")
        lines.append("")
        
        # Cost breakdown
        lines.append("### Cost Estimate")
        cost = estimates.cost_estimate
        lines.append(f"**Total Cost**: ${cost.total_cost:,.0f}")
        lines.append(f"- Human Resources: ${cost.human_resources:,.0f}")
        lines.append(f"- Infrastructure: ${cost.infrastructure:,.0f}")
        lines.append(f"- Software Licenses: ${cost.software_licenses:,.0f}")
        lines.append(f"- Contingency: ${cost.contingency:,.0f}")
        lines.append("")
        
        # Timeline
        lines.append("### Timeline")
        timeline = estimates.timeline_estimate
        lines.append(f"**Total Duration**: {timeline.total_duration_days} days ({timeline.total_duration_months:.1f} months)")
        lines.append(f"**Critical Path**: {timeline.critical_path_duration} days")
        lines.append(f"**Buffer**: {timeline.buffer_days} days")
        lines.append("")
        
        return lines

    async def _generate_html_for_pdf(self, document: Any, options: ExportOptions) -> str:
        """Generate HTML content optimized for PDF generation."""
        template_content = self.html_templates.get("pdf", self.html_templates["professional"])
        template = self.jinja_env.from_string(template_content)
        
        template_data = {
            "title": document.title,
            "document_type": document.document_type.replace('_', ' ').title(),
            "content": document.content,
            "metadata": document.metadata if options.include_metadata else {},
            "wbs": document.wbs if options.include_wbs else None,
            "estimates": document.estimates if options.include_estimates else None,
            "generated_at": document.created_at.strftime("%Y-%m-%d %H:%M:%S"),
            "page_size": options.page_size,
            "margins": options.margins
        }
        
        return template.render(**template_data)

    def _add_content_to_word(self, doc: Document, content: Dict[str, Any]) -> None:
        """Add content sections to Word document."""
        if content.get("overview"):
            doc.add_heading('Overview', level=1)
            doc.add_paragraph(str(content["overview"]))
        
        if content.get("requirements"):
            doc.add_heading('Requirements', level=1)
            requirements = content["requirements"]
            if isinstance(requirements, list):
                for req in requirements:
                    p = doc.add_paragraph(str(req), style='List Bullet')
            else:
                doc.add_paragraph(str(requirements))
        
        if content.get("implementation"):
            doc.add_heading('Implementation', level=1)
            impl = content["implementation"]
            if isinstance(impl, dict):
                if impl.get("approach"):
                    doc.add_paragraph(f"Approach: {impl['approach']}")
                if impl.get("technologies"):
                    doc.add_paragraph("Technologies:")
                    for tech in impl["technologies"]:
                        doc.add_paragraph(tech, style='List Bullet')
            else:
                doc.add_paragraph(str(impl))
        
        if content.get("timeline"):
            doc.add_heading('Timeline', level=1)
            timeline = content["timeline"]
            if isinstance(timeline, dict):
                if timeline.get("total_duration"):
                    doc.add_paragraph(f"Duration: {timeline['total_duration']}")
                if timeline.get("milestones"):
                    doc.add_paragraph("Milestones:")
                    for milestone in timeline["milestones"]:
                        if isinstance(milestone, dict):
                            doc.add_paragraph(f"{milestone.get('name', 'Milestone')} - {milestone.get('description', '')}", style='List Bullet')
            else:
                doc.add_paragraph(str(timeline))

    def _add_metadata_to_word(self, doc: Document, document: Any) -> None:
        """Add metadata section to Word document."""
        doc.add_paragraph(f"Document Type: {document.document_type}")
        doc.add_paragraph(f"Generated: {document.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Generation Time: {document.generation_time_ms}ms")

    def _add_wbs_to_word(self, doc: Document, wbs: Any) -> None:
        """Add WBS section to Word document."""
        doc.add_paragraph(f"Total Hours: {wbs.total_estimated_hours:,.1f}")
        doc.add_paragraph(f"Total Days: {wbs.total_estimated_days}")
        
        for phase in wbs.phases:
            doc.add_heading(phase.name, level=2)
            doc.add_paragraph(phase.description)
            doc.add_paragraph(f"Duration: {phase.estimated_duration_days} days")
            
            if phase.tasks:
                doc.add_paragraph("Tasks:", style='Heading 3')
                for task in phase.tasks:
                    doc.add_paragraph(f"{task.name} ({task.estimated_hours}h)", style='List Bullet')
                    doc.add_paragraph(f"  {task.description}")

    def _add_estimates_to_word(self, doc: Document, estimates: Any) -> None:
        """Add estimates section to Word document."""
        doc.add_heading('Team Composition', level=2)
        doc.add_paragraph(f"Team Size: {estimates.team_composition.total_team_size}")
        doc.add_paragraph(f"Monthly Cost: ${estimates.team_composition.estimated_monthly_cost:,.0f}")
        
        doc.add_heading('Cost Breakdown', level=2)
        cost = estimates.cost_estimate
        doc.add_paragraph(f"Total Cost: ${cost.total_cost:,.0f}")
        doc.add_paragraph(f"Human Resources: ${cost.human_resources:,.0f}")
        doc.add_paragraph(f"Infrastructure: ${cost.infrastructure:,.0f}")

    def _get_pdf_styles(self, options: ExportOptions) -> str:
        """Get CSS styles for PDF generation."""
        return f"""
        @page {{
            size: {options.page_size};
            margin: {options.margins['top']}in {options.margins['right']}in {options.margins['bottom']}in {options.margins['left']}in;
        }}
        body {{
            font-family: Arial, sans-serif;
            font-size: 11pt;
            line-height: 1.4;
            color: #333;
        }}
        h1 {{
            font-size: 18pt;
            font-weight: bold;
            color: #2c3e50;
            margin-bottom: 20px;
            border-bottom: 2px solid #3498db;
            padding-bottom: 5px;
        }}
        h2 {{
            font-size: 14pt;
            font-weight: bold;
            color: #2c3e50;
            margin-top: 20px;
            margin-bottom: 10px;
        }}
        h3 {{
            font-size: 12pt;
            font-weight: bold;
            color: #34495e;
            margin-top: 15px;
            margin-bottom: 8px;
        }}
        p {{
            margin-bottom: 8px;
            text-align: justify;
        }}
        ul {{
            margin-bottom: 12px;
        }}
        li {{
            margin-bottom: 4px;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin-bottom: 15px;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }}
        th {{
            background-color: #f2f2f2;
            font-weight: bold;
        }}
        """

    def _get_professional_template(self) -> str:
        """Get professional HTML template."""
        return """
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>{{ title }}</title>
            <style>
                body {
                    font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                    line-height: 1.6;
                    color: #333;
                    max-width: 800px;
                    margin: 0 auto;
                    padding: 20px;
                    background-color: #f9f9f9;
                }
                .container {
                    background-color: white;
                    padding: 40px;
                    border-radius: 8px;
                    box-shadow: 0 2px 10px rgba(0,0,0,0.1);
                }
                h1 {
                    color: #2c3e50;
                    border-bottom: 3px solid #3498db;
                    padding-bottom: 10px;
                    margin-bottom: 30px;
                }
                h2 {
                    color: #34495e;
                    border-left: 4px solid #3498db;
                    padding-left: 15px;
                    margin-top: 30px;
                }
                .metadata {
                    background-color: #ecf0f1;
                    padding: 15px;
                    border-radius: 5px;
                    margin-bottom: 20px;
                }
                .requirement {
                    background-color: #e8f5e8;
                    padding: 10px;
                    margin: 5px 0;
                    border-left: 4px solid #27ae60;
                }
            </style>
        </head>
        <body>
            <div class="container">
                <h1>{{ title }}</h1>
                
                {% if metadata %}
                <div class="metadata">
                    <strong>Document Type:</strong> {{ document_type }}<br>
                    <strong>Generated:</strong> {{ generated_at }}
                </div>
                {% endif %}
                
                {% if content.overview %}
                <h2>Overview</h2>
                <p>{{ content.overview }}</p>
                {% endif %}
                
                {% if content.requirements %}
                <h2>Requirements</h2>
                {% for req in content.requirements %}
                <div class="requirement">{{ req }}</div>
                {% endfor %}
                {% endif %}
                
                {% if content.implementation %}
                <h2>Implementation</h2>
                <p><strong>Approach:</strong> {{ content.implementation.approach }}</p>
                {% if content.implementation.technologies %}
                <p><strong>Technologies:</strong></p>
                <ul>
                {% for tech in content.implementation.technologies %}
                <li>{{ tech }}</li>
                {% endfor %}
                </ul>
                {% endif %}
                {% endif %}
            </div>
        </body>
        </html>
        """

    def _get_minimal_template(self) -> str:
        """Get minimal HTML template."""
        return """
        <!DOCTYPE html>
        <html>
        <head>
            <title>{{ title }}</title>
            <style>
                body { font-family: Arial, sans-serif; margin: 40px; }
                h1 { color: #333; }
                h2 { color: #666; margin-top: 25px; }
            </style>
        </head>
        <body>
            <h1>{{ title }}</h1>
            {% if content.overview %}
            <h2>Overview</h2>
            <p>{{ content.overview }}</p>
            {% endif %}
        </body>
        </html>
        """

    def _get_corporate_template(self) -> str:
        """Get corporate HTML template."""
        return self._get_professional_template()  # Use professional for now